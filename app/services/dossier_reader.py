"""Dossier Reader.

Reads every supported file in the user-supplied local folder and returns
a list of extracted chunks. Original files never leave the machine and
raw extracted text is never logged.

Supported file types and their handlers:

* ``.pdf``  -> ``pdfplumber`` (one chunk per page).
* ``.docx`` -> ``python-docx`` (paragraphs + tables, single chunk per file).
* ``.txt`` / ``.md`` -> read as utf-8 text.
* ``.json`` -> parsed via ``json``. The structured payload is preserved.
* ``.csv`` -> read as text plus a small ``json_data`` preview.
* ``.png`` / ``.jpg`` / ``.jpeg`` -> metadata only (no OCR yet).

Failures in a single file do not stop processing. The bad file is marked
``failed`` and reading continues.
"""

from __future__ import annotations

import csv
import hashlib
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional

from app.models.schemas import (
    SOURCE_PRIORITY,
    ChunkRecord,
    ExtractionStatus,
    SourceType,
)
from app.services.folder_validator import (
    IMAGE_EXTENSIONS,
    SUPPORTED_EXTENSIONS,
    TEXT_LIKE_EXTENSIONS,
    classify_source,
)
from app.utils.file_utils import (
    MAX_DOSSIER_FILE_BYTES,
    iter_contained_files,
    is_within_size_limit,
)


_PDF_MAX_PAGES = 100
_WARNING_MAX_CHARS = 200


def _chunk_id(file_path: str, suffix: str) -> str:
    digest = hashlib.sha1(file_path.encode("utf-8")).hexdigest()[:10]
    return f"{digest}::{suffix}"


def _sanitize_warning(message: Optional[str]) -> Optional[str]:
    """Collapse a warning to a single safe line.

    Reader warnings only ever carry exception class names and short
    operational notes — never raw dossier text — but we still flatten
    newlines and cap the length so nothing large or multi-line can leak
    into a log line or the debug panel.
    """
    if not message:
        return None
    flat = " ".join(str(message).split())
    if len(flat) > _WARNING_MAX_CHARS:
        flat = flat[: _WARNING_MAX_CHARS - 1].rstrip() + "…"
    return flat or None


def _detect_section(text: str) -> Optional[str]:
    """Best-effort section name from a Markdown heading, if present."""
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            heading = stripped.lstrip("#").strip()
            if heading:
                return heading[:120]
        if stripped:
            # Only the first non-blank line can be the document heading.
            break
    return None


def _load_json(path: Path) -> tuple[Optional[object], Optional[str]]:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as fh:
            return json.load(fh), None
    except (OSError, ValueError) as exc:
        return None, f"{exc.__class__.__name__}: {exc}"


def _read_text(path: Path) -> tuple[str, Optional[str], Optional[str]]:
    """Return ``(text, fatal_error, decode_warning)``.

    ``fatal_error`` is set only when the file could not be opened/read at
    all. ``decode_warning`` is set (but the text is still usable) when some
    non-UTF-8 bytes had to be dropped — so a mis-encoded file degrades
    visibly instead of producing quietly-corrupted evidence text.
    """
    try:
        raw = path.read_bytes()
    except OSError as exc:
        return "", f"{exc.__class__.__name__}: {exc}", None
    text = raw.decode("utf-8", errors="ignore")
    try:
        raw.decode("utf-8")
        warning = None
    except UnicodeDecodeError:
        warning = "some non-UTF-8 bytes were dropped while reading"
    return text, None, warning


def _read_csv_preview(path: Path, max_rows: int = 50) -> tuple[str, list[list[str]], Optional[str]]:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore", newline="") as fh:
            reader = csv.reader(fh)
            rows: list[list[str]] = []
            for i, row in enumerate(reader):
                if i >= max_rows:
                    break
                rows.append(row)
    except OSError as exc:
        return "", [], f"{exc.__class__.__name__}: {exc}"

    text_lines = ["\t".join(cell for cell in row) for row in rows]
    return "\n".join(text_lines), rows, None


def _pdf_chunks(
    path: Path,
    base_meta: dict,
) -> tuple[list[ChunkRecord], Optional[str]]:
    try:
        import pdfplumber  # type: ignore
    except Exception as exc:  # pragma: no cover - dependency must be present
        return [], f"pdfplumber unavailable: {exc}"

    records: list[ChunkRecord] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            page_iter = pdf.pages[:_PDF_MAX_PAGES]
            for index, page in enumerate(page_iter, start=1):
                try:
                    text = page.extract_text() or ""
                except Exception as exc:  # noqa: BLE001 - per-page tolerance
                    records.append(
                        ChunkRecord(
                            chunk_id=_chunk_id(base_meta["file_path"], f"p{index}"),
                            page_number=index,
                            extracted_text="",
                            extraction_status="failed",
                            extraction_warning=f"page extract failed: {exc.__class__.__name__}",
                            **base_meta,
                        )
                    )
                    continue
                records.append(
                    ChunkRecord(
                        chunk_id=_chunk_id(base_meta["file_path"], f"p{index}"),
                        page_number=index,
                        extracted_text=text.strip(),
                        extraction_status="ok" if text.strip() else "empty",
                        **base_meta,
                    )
                )
    except Exception as exc:  # noqa: BLE001 - one bad PDF must not crash run
        return records, f"pdfplumber failed: {exc.__class__.__name__}"
    return records, None


def _docx_chunk(
    path: Path,
    base_meta: dict,
) -> tuple[list[ChunkRecord], Optional[str]]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover
        return [], f"python-docx unavailable: {exc}"

    try:
        document = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        return [], f"docx open failed: {exc.__class__.__name__}"

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    table_texts: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_texts.append(" | ".join(cells))

    full_text = "\n".join(paragraphs + table_texts).strip()
    status: ExtractionStatus = "ok" if full_text else "empty"
    return (
        [
            ChunkRecord(
                chunk_id=_chunk_id(base_meta["file_path"], "docx"),
                extracted_text=full_text,
                extraction_status=status,
                **base_meta,
            )
        ],
        None,
    )


def _failed_chunk(path: Path, base_meta: dict, message: str) -> ChunkRecord:
    return ChunkRecord(
        chunk_id=_chunk_id(base_meta["file_path"], "failed"),
        extracted_text="",
        extraction_status="failed",
        extraction_warning=_sanitize_warning(message) or "extraction failed",
        **base_meta,
    )


def _build_base_meta(
    path: Path,
    rel_path: Path,
    source_type: SourceType,
) -> dict:
    return {
        "file_name": path.name,
        "file_path": str(rel_path),
        "file_type": path.suffix.lower(),
        "source_type": source_type,
        "source_priority": SOURCE_PRIORITY[source_type],
    }


def _fallback_base_meta(path: Path, rel_path: Path) -> dict:
    """Base metadata for a file whose classification itself failed."""
    return _build_base_meta(path, rel_path, "unknown_supported_file")


def read_dossier(folder_path: str | Path) -> list[ChunkRecord]:
    """Walk the dossier folder and return extracted chunks.

    One file may produce multiple chunks (e.g., one per PDF page). Files
    that fail to open produce a single ``failed`` chunk with the warning
    captured in ``extraction_warning``. Unsupported files are skipped.
    """
    folder = Path(folder_path).expanduser()
    if not folder.exists() or not folder.is_dir():
        return []

    chunks: list[ChunkRecord] = []
    try:
        # Symlink-safe, containment-checked, count-bounded walk. A symlink
        # whose target lives outside the folder is never read, so files the
        # user did not put in the dossier can never reach the LLM.
        candidates = list(iter_contained_files(folder))
    except OSError:
        return []

    for path in candidates:
        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue

        rel_path = path.relative_to(folder)

        # Skip files larger than the per-file cap rather than reading them
        # fully into memory; record a visible failed chunk so the UI can
        # show the file was intentionally skipped.
        if not is_within_size_limit(path):
            base_meta = _fallback_base_meta(path, rel_path)
            mb = MAX_DOSSIER_FILE_BYTES // (1024 * 1024)
            chunks.append(
                _failed_chunk(
                    path,
                    base_meta,
                    f"file skipped: exceeds {mb}MB size limit",
                )
            )
            continue

        try:
            chunks.extend(_read_one_file(path, rel_path, ext))
        except Exception as exc:  # noqa: BLE001 - one bad file must not abort the run
            base_meta = _fallback_base_meta(path, rel_path)
            chunks.append(
                _failed_chunk(
                    path,
                    base_meta,
                    f"unexpected reader error: {exc.__class__.__name__}",
                )
            )

    return chunks


def _read_one_file(path: Path, rel_path: Path, ext: str) -> list[ChunkRecord]:
    """Extract one supported file into one or more chunks.

    Per-file failures are represented as ``failed`` chunks rather than
    raised exceptions; the caller wraps this in a final safety net so any
    truly unexpected error still degrades to a single failed chunk.
    """
    # Classification: read a small sample for JSON / text-like files so
    # the classifier can use content as a signal.
    text_sample = ""
    json_payload: object = None
    if ext in TEXT_LIKE_EXTENSIONS:
        try:
            with open(path, "rb") as fh:
                text_sample = fh.read(8192).decode("utf-8", errors="ignore")
        except OSError:
            text_sample = ""
        if ext == ".json":
            json_payload, _ = _load_json(path)

    source_type, _note = classify_source(
        path, text_sample=text_sample, json_payload=json_payload
    )
    base_meta = _build_base_meta(path, rel_path, source_type)

    if ext == ".pdf":
        pdf_chunks, warning = _pdf_chunks(path, base_meta)
        if pdf_chunks:
            if warning:
                pdf_chunks[-1].extraction_warning = _sanitize_warning(warning)
            return pdf_chunks
        return [_failed_chunk(path, base_meta, warning or "no pages")]

    if ext == ".docx":
        docx_chunks, warning = _docx_chunk(path, base_meta)
        if docx_chunks:
            return docx_chunks
        return [_failed_chunk(path, base_meta, warning or "no content")]

    if ext in {".txt", ".md"}:
        text, err, decode_warning = _read_text(path)
        status: ExtractionStatus = "ok" if text.strip() else "empty"
        section = _detect_section(text) if ext == ".md" else None
        return [
            ChunkRecord(
                chunk_id=_chunk_id(base_meta["file_path"], "text"),
                section_name=section,
                extracted_text=text.strip(),
                extraction_status="failed" if err else status,
                extraction_warning=_sanitize_warning(err or decode_warning),
                **base_meta,
            )
        ]

    if ext == ".json":
        payload, err = _load_json(path)
        if err:
            return [_failed_chunk(path, base_meta, err)]
        try:
            pretty = json.dumps(payload, indent=2, ensure_ascii=False)
        except Exception:  # noqa: BLE001
            pretty = ""
        return [
            ChunkRecord(
                chunk_id=_chunk_id(base_meta["file_path"], "json"),
                extracted_text=pretty,
                json_data=payload,
                extraction_status="ok" if payload is not None else "empty",
                **base_meta,
            )
        ]

    if ext == ".csv":
        text, rows, err = _read_csv_preview(path)
        return [
            ChunkRecord(
                chunk_id=_chunk_id(base_meta["file_path"], "csv"),
                extracted_text=text,
                json_data={"rows": rows} if rows else None,
                extraction_status="failed" if err else ("ok" if rows else "empty"),
                extraction_warning=_sanitize_warning(err),
                **base_meta,
            )
        ]

    if ext in IMAGE_EXTENSIONS:
        return [
            ChunkRecord(
                chunk_id=_chunk_id(base_meta["file_path"], "image"),
                extracted_text="",
                extraction_status="metadata_only",
                extraction_warning="image stored as metadata only; OCR not run",
                **base_meta,
            )
        ]

    return []


def chunks_by_file(chunks: Iterable[ChunkRecord]) -> dict[str, list[ChunkRecord]]:
    by_file: dict[str, list[ChunkRecord]] = {}
    for chunk in chunks:
        by_file.setdefault(chunk.file_path, []).append(chunk)
    return by_file


# ---------------------------------------------------------------------------
# Read summary (UI- and log-safe view over the extracted chunks)
# ---------------------------------------------------------------------------


@dataclass
class FileReadSummary:
    """Per-file rollup used by the UI summary and debug panel.

    Carries metadata and a sanitized warning only — never raw text.
    """

    file_name: str
    file_path: str
    file_type: str
    source_type: str
    chunk_count: int
    status: str
    warning: Optional[str] = None


@dataclass
class DossierReadSummary:
    """Aggregate, text-free summary of a dossier read.

    Safe to log and to render: it counts files and chunks and lists
    per-file status, but contains none of the extracted dossier text.
    """

    files_processed: int = 0
    chunks_extracted: int = 0
    failed_files: int = 0
    files: list[FileReadSummary] = field(default_factory=list)


def _aggregate_status(statuses: list[str]) -> str:
    """Reduce a file's per-chunk statuses to one headline status."""
    if not statuses:
        return "empty"
    if all(s == "failed" for s in statuses):
        return "failed"
    for preferred in ("ok", "partial", "metadata_only", "empty"):
        if preferred in statuses:
            return preferred
    return statuses[0]


def summarize_chunks(chunks: Iterable[ChunkRecord]) -> DossierReadSummary:
    """Build a text-free summary of an extracted dossier.

    ``files_processed`` counts the distinct files that produced at least
    one chunk, ``chunks_extracted`` is the total chunk count, and
    ``failed_files`` counts files whose every chunk failed to extract.
    """
    by_file = chunks_by_file(chunks)
    file_summaries: list[FileReadSummary] = []
    failed_files = 0
    total_chunks = 0

    for file_path, file_chunks in by_file.items():
        total_chunks += len(file_chunks)
        statuses = [c.extraction_status for c in file_chunks]
        status = _aggregate_status(statuses)
        if status == "failed":
            failed_files += 1
        warning = next(
            (c.extraction_warning for c in file_chunks if c.extraction_warning),
            None,
        )
        first = file_chunks[0]
        file_summaries.append(
            FileReadSummary(
                file_name=first.file_name,
                file_path=file_path,
                file_type=first.file_type,
                source_type=first.source_type,
                chunk_count=len(file_chunks),
                status=status,
                warning=_sanitize_warning(warning),
            )
        )

    return DossierReadSummary(
        files_processed=len(by_file),
        chunks_extracted=total_chunks,
        failed_files=failed_files,
        files=file_summaries,
    )
